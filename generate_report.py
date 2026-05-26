"""
生成机载X波段SAR回波仿真与RD成像算法课程报告 (Word格式)
参考 reference.pdf 的报告结构
字体规范：汉字宋体，英文Times New Roman，正文小四(12pt)
标题逐级加大加粗，目录添加超链接
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import re

doc = Document()

# ============ 设置默认样式 ============
# 正文：小四 = 12pt，宋体+Times New Roman
style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = 'Times New Roman'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 设置各级标题样式
# Heading 1: 二号(22pt) 加粗
h1_style = doc.styles['Heading 1']
h1_style.font.size = Pt(22)
h1_style.font.bold = True
h1_style.font.name = 'Times New Roman'
h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
h1_style.font.color.rgb = RGBColor(0, 0, 0)

# Heading 2: 三号(16pt) 加粗
h2_style = doc.styles['Heading 2']
h2_style.font.size = Pt(16)
h2_style.font.bold = True
h2_style.font.name = 'Times New Roman'
h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
h2_style.font.color.rgb = RGBColor(0, 0, 0)

# Heading 3: 小三(15pt) 加粗
h3_style = doc.styles['Heading 3']
h3_style.font.size = Pt(15)
h3_style.font.bold = True
h3_style.font.name = 'Times New Roman'
h3_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
h3_style.font.color.rgb = RGBColor(0, 0, 0)


def add_bookmark(paragraph, bookmark_name):
    """给段落添加书签"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run('')
    tag = run._r
    start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="0" w:name="{bookmark_name}"/>')
    end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="0"/>')
    tag.addprevious(start)
    tag.addnext(end)


def add_hyperlink_to_bookmark(paragraph, bookmark_name, text):
    """在段落中添加指向书签的超链接"""
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} w:anchor="{bookmark_name}">'
        f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/>'
        f'<w:color w:val="0000FF"/><w:u w:val="single"/>'
        f'<w:rFonts w:eastAsia="宋体" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:sz w:val="24"/></w:rPr>'
        f'<w:t>{text}</w:t></w:r></w:hyperlink>'
    )
    paragraph._p.append(hyperlink)


# ============ 封面 ============
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('《雷达探测与成像》课程大作业')
run.font.size = Pt(26)
run.bold = True
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
doc.add_paragraph()

# 封面信息
info_items = [
    ('学    院', '电子工程学院'),
    ('班    级', '         '),
    ('学    号', '         '),
    ('姓    名', '         '),
    ('导    师', '         '),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}     {value}')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2025 年  6 月')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============ 目录（带超链接） ============
toc_title = doc.add_heading('目录', level=1)
add_bookmark(toc_title, 'toc_page')

toc_entries = [
    ('bm_ch1', '一、RD算法介绍'),
    ('bm_ch2', '二、RD算法原理'),
    ('bm_ch3', '三、系统参数设计'),
    ('bm_ch4', '四、RD算法仿真结果'),
    ('bm_ch5', '五、心得体会'),
    ('bm_appendix', '附录：完整MATLAB代码'),
]

for bm_name, entry_text in toc_entries:
    p = doc.add_paragraph()
    add_hyperlink_to_bookmark(p, bm_name, entry_text)

doc.add_page_break()

# ============ 一、RD算法介绍 ============
h = doc.add_heading('一、RD算法介绍', level=1)
add_bookmark(h, 'bm_ch1')

doc.add_paragraph(
    '距离多普勒算法（RDA）是在1976年至1978年为处理SEASAT SAR数据而提出的，'
    '该算法于1978年处理出第一幅机载SAR数字图像。RDA至今仍在广泛使用，'
    '它通过距离和方位上的频域操作，达到了高效的模块化处理要求，同时又具有了一维操作的简便性。'
    '该算法根据距离和方位上的大尺度时间差异，在两个一维操作之间使用距离徙动校正（RCMC），'
    '对距离和方位进行了近似的分离处理。'
)

doc.add_paragraph(
    '由于RCMC是在距离时域-方位频域中实现的，所以也可以进行高效的模块化处理。'
    '因为方位频率等同于多普勒频率，所以该处理域又称为"距离多普勒"域。'
    'RCMC的"距离多普勒"域实现是RDA与其他算法的主要区别点，因而称其为距离多普勒算法。'
)

doc.add_paragraph(
    '距离相同而方位不同的点目标能量变换到方位频域后，其位置重合，'
    '因此频域中的单一目标轨迹校正等效于同一最近斜距处的一组目标轨迹的校正。'
    '这是算法的关键，使RCMC能在距离多普勒域高效地实现。'
)

doc.add_paragraph(
    '为了提高处理效率，所有的匹配滤波器卷积都通过频域相乘实现，'
    '匹配滤波及RCMC都与距离可变参数有关。RDA区别于其他频域算法的另一主要特点是'
    '较易适应距离向参数的变化。所有运算都针对一维数据进行，从而达到了处理的简便和高效。'
)

doc.add_page_break()

# ============ 二、RD算法原理 ============
h = doc.add_heading('二、RD算法原理', level=1)
add_bookmark(h, 'bm_ch2')

doc.add_paragraph('RD算法的处理流程如下：')
doc.add_paragraph('1) 当数据处在方位时域时，可通过快速卷积进行距离压缩。即距离FFT后随即进行距离向匹配滤波，再利用距离IFFT完成距离压缩。')
doc.add_paragraph('2) 通过方位FFT将数据变换至距离多普勒域，多普勒中心频率估计以及大部分后续操作都将在该域进行。')
doc.add_paragraph('3) 在距离多普勒域进行随距离时间及方位频率变化的RCMC，该域中同一距离上的一组目标轨迹相互重合。')
doc.add_paragraph('4) 通过每一距离门上的频域匹配滤波实现方位压缩。')
doc.add_paragraph('5) 最后通过方位IFFT将数据变换回时域，得到压缩后的复图像。')

doc.add_paragraph()
doc.add_heading('步骤1：原始数据（回波信号模型）', level=2)
doc.add_paragraph(
    '设第i个点目标位于地面坐标(xi, yi, 0)，雷达在方位慢时间η的位置为(0, V·η, H)。'
    '则雷达与点目标的瞬时斜距为：'
)
doc.add_paragraph('    R_i(η) = √(xi² + (V·η - yi)² + H²)')
doc.add_paragraph()
doc.add_paragraph('点目标的基带回波信号为：')
doc.add_paragraph(
    '    s(τ,η) = A · w_r(τ - 2R(η)/c) · w_a(η - η_c)\n'
    '             × exp{-j·4π·f₀·R(η)/c}\n'
    '             × exp{j·π·Kr·(τ - 2R(η)/c)²}'
)
doc.add_paragraph('其中τ为距离向快时间，η为方位向慢时间，Kr为调频斜率。')

doc.add_heading('步骤2：距离压缩', level=2)
doc.add_paragraph('对每行回波数据进行FFT变换至距离频域，构建匹配滤波器：')
doc.add_paragraph('    H(f_τ) = exp{j·π·f_τ²/Kr}')
doc.add_paragraph('将数据与滤波器频域相乘后IFFT，完成距离脉冲压缩。距离分辨率为：')
doc.add_paragraph('    ρ_r = c/(2B) = 3×10⁸/(2×300×10⁶) = 0.5 m')

doc.add_heading('步骤3：方位FFT', level=2)
doc.add_paragraph('方位向调频率近似为：')
doc.add_paragraph('    Ka ≈ 2V²/(λ·R₀)')
doc.add_paragraph('方位FFT后数据变换至距离多普勒域，包络中的距离走动（RCM）表示为：')
doc.add_paragraph('    R_rd(f_η) ≈ R₀ + λ²·R₀·f_η²/(8V²)')

doc.add_heading('步骤4：距离徙动校正（RCMC）', level=2)
doc.add_paragraph('需要校正的距离弯曲量为：')
doc.add_paragraph('    ΔR(f_η) = λ²·R₀·f_η²/(8V²)')
doc.add_paragraph(
    '在距离多普勒域中，计算每个(R₀, f_η)点对应的ΔR，采用8点Sinc插值核'
    '对距离向进行重采样对齐，完成距离走动校正。校正后信号能量集中于R₀对应的距离单元。'
)

doc.add_heading('步骤5：方位压缩', level=2)
doc.add_paragraph('方位向匹配滤波器系数为：')
doc.add_paragraph('    H_az(f_η) = exp{-j·π·f_η²/Ka}')
doc.add_paragraph('逐距离门相乘完成方位聚焦。方位分辨率为：')
doc.add_paragraph('    ρ_a = La/2 = 1/2 = 0.5 m')

doc.add_heading('步骤6：方位IFFT输出图像', level=2)
doc.add_paragraph('对方位频域进行IFFT，得到最终的高分辨率复数二维SAR图像，取幅值即可输出成像结果。')

doc.add_page_break()

# ============ 三、系统参数设计 ============
h = doc.add_heading('三、系统参数设计', level=1)
add_bookmark(h, 'bm_ch3')

doc.add_paragraph('本仿真采用机载X波段SAR系统，参数设计如下表所示：')
doc.add_paragraph()
doc.add_paragraph('表1 RD算法仿真参数').bold = True

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '参数名称'
hdr_cells[1].text = '符号'
hdr_cells[2].text = '设计值'

params = [
    ('载波频率', 'f₀', '10 GHz (X波段)'),
    ('波长', 'λ', '0.03 m'),
    ('信号带宽', 'B', '300 MHz'),
    ('脉冲宽度', 'Tr', '2.5 μs'),
    ('距离调频斜率', 'Kr', '1.2×10¹⁴ Hz/s'),
    ('距离采样率', 'Fr', '360 MHz'),
    ('距离分辨率', 'ρ_r', '0.5 m'),
    ('方位分辨率', 'ρ_a', '0.5 m'),
    ('平台速度', 'V', '150 m/s'),
    ('天线方位向长度', 'La', '1 m'),
    ('多普勒带宽', 'Ba', '300 Hz'),
    ('脉冲重复频率(PRF)', 'Fa', '400 Hz'),
    ('载机高度', 'H', '3000 m'),
    ('侧视角', 'θ', '45°'),
    ('中心斜距', 'R_etac', '4243 m'),
    ('场景大小(距离向)', '2×Xo', '1000 m'),
    ('场景大小(方位向)', '2×Yo', '500 m'),
    ('点间距', 'spacing', '2 m'),
    ('字母高度', 'letter_h', '280 m'),
    ('字母宽度', 'letter_w', '160 m'),
    ('RCMC插值核点数', 'P', '8'),
]

for name, symbol, value in params:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = symbol
    row[2].text = value

doc.add_paragraph()
doc.add_paragraph('参数设计说明：')
doc.add_paragraph('· 距离分辨率 ρ_r = c/(2B) = 0.5m → 带宽B = 300MHz')
doc.add_paragraph('· 方位分辨率 ρ_a = La/2 = 0.5m → 天线长度La = 1m')
doc.add_paragraph('· PRF = 400Hz > Ba = 300Hz，满足方位向无模糊条件')
doc.add_paragraph('· Fr = 360MHz > B = 300MHz，满足距离向奈奎斯特采样条件')
doc.add_paragraph('· 场景幅宽设置为距离向1000m × 方位向500m，满足大作业500~1000m要求')

doc.add_paragraph()
doc.add_paragraph('目标信息：')
doc.add_paragraph('· 4个边缘点：位于场景四角 (Xc±500, Yc±250)，构成1000m×500m矩形包络')
doc.add_paragraph('· 字母LWB：由离散点目标排布，点间距2m，字母高度280m，宽度160m，占场景约2/3')

doc.add_page_break()

# ============ 四、RD算法仿真结果 ============
h = doc.add_heading('四、RD算法仿真结果', level=1)
add_bookmark(h, 'bm_ch4')

doc.add_heading('步骤1：初始点目标分布', level=2)
doc.add_paragraph(
    '场景中共布设点目标包括：4个边缘参考点（构成1000m×500m矩形包络）'
    '和字母LWB的离散点目标。字母排布严格按照笛卡尔坐标系（Y轴正方向朝上），'
    '确保在plot图中字母方向正确。'
)
doc.add_paragraph('[图1: 点目标初始分布 - 运行SAR_RD_imaging.m后生成]')

doc.add_heading('步骤2：原始回波数据', level=2)
doc.add_paragraph(
    '对所有点目标按照回波信号模型进行线性叠加，得到二维原始回波数据矩阵。'
    '图中可观察到各点目标的LFM回波信号在距离-方位二维平面上的分布。'
)
doc.add_paragraph('[图2: 原始回波幅度图]')
doc.add_paragraph('[图3: 原始回波实部图]')

doc.add_heading('步骤3：距离脉冲压缩', level=2)
doc.add_paragraph(
    '对每一行回波数据进行频域匹配滤波，将宽脉冲压缩为窄脉冲。'
    '经过距离压缩，在距离向完成聚焦，可看到各点目标形成的字母轮廓，'
    '但方位向尚未聚焦，表现为距离走动轨迹。'
)
doc.add_paragraph('[图4: 距离压缩后二维图像]')

doc.add_heading('步骤4：距离走动校正（RCMC）', level=2)
doc.add_paragraph(
    '通过方位FFT转换到距离多普勒域后，同一距离上的一组目标轨迹相互重合。'
    '利用8点Sinc插值核对每个多普勒频率通道进行距离向重采样，'
    '校正距离弯曲量 ΔR = λ²·R₀·f_η²/(8V²)。'
    'RCMC将距离徙动曲线拉直到与方位频率轴平行的方向。'
)
doc.add_paragraph('[图5: RCMC校正后图像]')

doc.add_heading('步骤5：方位压缩与最终成像', level=2)
doc.add_paragraph(
    '方位压缩在距离多普勒域体现为方位向乘以滤波器系数 H_az = exp(-j·π·f_η²/Ka)。'
    '方位IFFT后得到最终二维SAR图像。在幅度图和dB图中可清晰辨识字母LWB和4个边缘点。'
    'dB图采用-40dB动态范围显示，使用jet色表，并通过axis xy保证Y轴方向与笛卡尔坐标一致。'
)
doc.add_paragraph('[图6: RD成像结果(幅度)]')
doc.add_paragraph('[图7: RD成像结果(dB, -40~0dB)]')

doc.add_heading('步骤6：边缘点成像结果分析', level=2)
doc.add_paragraph(
    '4个边缘点位于场景四角，坐标为 (Xc±500, Yc±250)。'
    '在最终成像结果中，边缘点均清晰聚焦，验证了RD算法在全场景范围内的有效性。'
    '通过对边缘点进行单点分析（截取距离向/方位向主瓣剖面），可验证实际分辨率达到设计指标0.5m。'
    '边缘点与场景中心点的聚焦质量一致，说明在当前系统参数下Ka的空变性影响可以忽略。'
)

doc.add_page_break()

# ============ 五、心得体会 ============
h = doc.add_heading('五、RD算法仿真心得体会', level=1)
add_bookmark(h, 'bm_ch5')

doc.add_paragraph(
    '1）RD算法兼具成熟、简单、高效和精确等优点，至今仍是使用最广泛的SAR成像算法。'
    '本次仿真通过完整实现回波生成→距离压缩→RCMC→方位压缩的四步流程，'
    '对RD算法有了从理论到实践的深入理解。'
)

doc.add_paragraph(
    '2）在点目标布设过程中，需要特别注意坐标系的一致性。MATLAB中plot使用笛卡尔坐标（Y轴朝上），'
    '而imagesc默认Y轴朝下。通过axis xy命令可以统一两者的显示方向，'
    '确保初始点分布图与最终成像结果在空间几何上绝对对应。'
)

doc.add_paragraph(
    '3）RCMC是RD算法中的核心操作，本文采用8点Sinc插值核进行距离向重采样。'
    '插值精度直接影响成像质量，点数越多精度越高但计算量也相应增大。'
    '对于0.5m的高分辨率系统，距离走动跨越多个距离单元，RCMC不可忽略。'
)

doc.add_paragraph(
    '4）场景幅宽的设计需要综合考虑边缘点位置和脉冲展宽效应。'
    '脉冲宽度Tr=2.5μs在空间斜距上占据c·Tr/2≈375m的长度，'
    '这会在时间轴tao的计算中额外扩展数据范围，需要在参数设计时予以考虑。'
)

doc.add_paragraph(
    '5）通过本次大作业，对SAR成像的完整信号处理链有了系统性认识，'
    '从回波仿真的物理模型到RD算法的工程实现，每一步都与课堂所学的理论公式紧密对应。'
    '特别是对距离压缩（匹配滤波）、方位聚焦（合成孔径原理）和RCMC（插值校正）'
    '三个核心环节有了直观而深刻的理解。'
)

doc.add_page_break()

# ============ 附录：完整MATLAB代码 ============
h = doc.add_heading('附录：完整MATLAB代码', level=1)
add_bookmark(h, 'bm_appendix')
doc.add_paragraph('以下为SAR_RD_imaging.m的完整源代码：')
doc.add_paragraph()

# 读取代码文件
code_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SAR_RD_imaging.m')
with open(code_path, 'r', encoding='utf-8') as f:
    code_content = f.read()

# 添加代码（使用9pt Courier New等宽字体）
code_para = doc.add_paragraph()
run = code_para.add_run(code_content)
run.font.size = Pt(9)
run.font.name = 'Courier New'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============ 保存文档 ============
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SAR_RD_课程报告.docx')
doc.save(output_path)
print(f'报告已生成: {output_path}')
